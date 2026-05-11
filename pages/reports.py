"""Octa Project Risk Log — Export Reports."""
import streamlit as st
import io
from datetime import date

from modules.auth import require_auth
from modules.sso import auto_login_from_url
from modules.ui_helpers import (inject_css, sidebar_nav, page_header,
                                 section_label, DARK)
from modules.database import (get_project, get_risks, get_risk_actions,
                               get_risk_reviews, get_risk_wps, get_work_packages)
from modules.charts import (chart_risk_matrix, chart_risk_levels,
                             chart_risk_status, chart_risk_by_category,
                             chart_risk_trend, chart_action_status)
from config import (DARK as D, CAT_LABELS, LEVEL_COLORS,
                    RISK_STATUS_OPTS, ACTION_TYPES)

st.set_page_config(page_title="Reports — Octa", page_icon="📥",
                   layout="wide", initial_sidebar_state="expanded")
inject_css(); auto_login_from_url(); require_auth(); sidebar_nav()

sel_pid = st.session_state.get("selected_project_id","")
if not sel_pid: st.switch_page("app.py"); st.stop()

proj    = get_project(sel_pid)
acronym = proj.get("acronym","") or sel_pid
page_header("Export Risk Reports",
            f"{acronym} — Download risk register as Word or HTML dashboard", "📥")
if st.button("← Dashboard"): st.switch_page("app.py")

muted  = D["muted"]
risks   = get_risks(sel_pid)
actions = get_risk_actions(sel_pid)
reviews = get_risk_reviews(sel_pid)
wps     = get_work_packages(sel_pid)
wp_map  = {w["wp_id"]: w.get("wp_number","") for w in wps}


# ── Word export ───────────────────────────────────────────────────────────────
def _build_word() -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2.5)
        section.page_width  = Cm(29.7)   # A4 landscape
        section.page_height = Cm(21)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(9)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Project Risk Log — {acronym}")
    run.bold = True; run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1B,0x2A,0x4A)
    doc.add_paragraph(f"Generated: {date.today().isoformat()}").runs[0].font.color.rgb = RGBColor(0x88,0x99,0xB0)
    doc.add_paragraph()

    # Risk register table
    doc.add_heading("Risk Register", level=1)
    LEVEL_BG = {"critical":"FC8181","high":"F6AD55","medium":"F6CC52","low":"6FCF97"}
    headers  = ["ID","Title","Category","Likelihood","Severity","Level",
                "Related WPs","Mitigation Strategy","Contingency","Status","Escalation"]
    widths   = [1.2,3.5,2.5,1.8,1.8,1.8,2.0,4.5,3.5,2.0,2.0]

    from docx.oxml.ns import qn
    from docx.oxml    import OxmlElement

    def _set_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _set_bg(hdr[i], "1B2A4A")

    level_order = {"critical":0,"high":1,"medium":2,"low":3}
    for r in sorted(risks, key=lambda x: level_order.get(x.get("risk_level","low"),3)):
        level = r.get("risk_level","medium")
        linked= get_risk_wps(r["risk_id"])
        wp_str= ", ".join(wp_map.get(wid,"") for wid in linked if wid in wp_map)
        row   = table.add_row().cells
        vals  = [
            r.get("risk_number",""),
            r.get("risk_title",""),
            CAT_LABELS.get(r.get("risk_category",""),r.get("risk_category","")),
            r.get("likelihood","").title(),
            r.get("severity","").title(),
            level.title(),
            wp_str,
            r.get("mitigation_strategy",""),
            r.get("contingency_plan",""),
            r.get("status","").title(),
            r.get("escalation_level","").title(),
        ]
        for i,val in enumerate(vals):
            row[i].text = val
        _set_bg(row[5], LEVEL_BG.get(level,"FFFFFF"))
        row[5].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Mitigation Actions", level=1)
    act_headers = ["Risk","Action","Type","Responsible","Due Date","Status","Progress","Notes"]
    act_table   = doc.add_table(rows=1, cols=len(act_headers))
    act_table.style = "Table Grid"
    ahdr = act_table.rows[0].cells
    for i,h in enumerate(act_headers):
        ahdr[i].text = h
        ahdr[i].paragraphs[0].runs[0].bold = True
        ahdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _set_bg(ahdr[i],"2E4A7A")

    risk_map2 = {r["risk_id"]: r for r in risks}
    ACTION_TYPE_LABELS = dict(ACTION_TYPES)
    for a in actions:
        r_obj = risk_map2.get(a.get("risk_id"),{})
        rnum  = r_obj.get("risk_number","")
        row   = act_table.add_row().cells
        vals  = [
            f"{rnum}: {r_obj.get('risk_title','')[:30]}",
            a.get("action_title",""),
            ACTION_TYPE_LABELS.get(a.get("action_type",""),""),
            a.get("responsible_person",""),
            a.get("due_date",""),
            a.get("status","").title(),
            f"{a.get('progress_pct',0)}%",
            a.get("outcome_notes",""),
        ]
        for i,val in enumerate(vals):
            row[i].text = val

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.read()


# ── HTML dashboard ────────────────────────────────────────────────────────────
def _build_html() -> str:
    def _c(fig): return fig.to_html(full_html=False, include_plotlyjs=False)
    return f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{acronym} — Risk Dashboard</title>
<script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
<style>
body{{background:#0f1421;color:#e2e8f0;font-family:Calibri,sans-serif;margin:0;padding:1rem 2rem}}
h1{{color:#00BCD4;font-size:1.8rem;margin-bottom:.3rem}}
p{{color:#8899b0;font-size:.9rem}}
.grid{{display:grid;grid-template-columns:1fr 1fr 1fr;gap:1.5rem;margin-top:1.5rem}}
.chart{{background:#1a2235;border-radius:12px;padding:1rem;border:1px solid rgba(255,255,255,.09)}}
.full{{grid-column:1/-1}}
footer{{margin-top:2rem;color:#8899b0;font-size:.78rem;text-align:center}}
</style></head>
<body>
<h1>🛡️ {acronym} — Project Risk Dashboard</h1>
<p>{proj.get('proposal_title','')} · Generated {date.today().isoformat()}</p>
<div class="grid">
  <div class="chart">{_c(chart_risk_matrix(risks,'Risk Matrix'))}</div>
  <div class="chart">{_c(chart_risk_levels(risks,'Risks by Level'))}</div>
  <div class="chart">{_c(chart_risk_status(risks,'Risks by Status'))}</div>
  <div class="chart full">{_c(chart_risk_by_category(risks,'Risks by Category'))}</div>
  {'<div class="chart full">' + _c(chart_risk_trend(reviews,'Risk Trend')) + '</div>' if reviews else ''}
  {'<div class="chart">' + _c(chart_action_status(actions,'Action Status')) + '</div>' if actions else ''}
</div>
<footer>Generated by Octa Platform · Project Risk Log · {acronym}</footer>
</body></html>"""


# ── UI ────────────────────────────────────────────────────────────────────────
section_label("📄 Word Report — Risk Register + Actions")
st.markdown(
    f"<p style='color:{muted};font-size:0.85rem'>"
    f"Downloads a Word document with the full risk register table and mitigation "
    f"actions table in EU reporting format.</p>", unsafe_allow_html=True)

if st.button("🔨 Generate Word Report", type="primary"):
    with st.spinner("Building risk register…"):
        docx_bytes = _build_word()
    st.download_button(
        f"📥 Download {acronym}_Risk_Register.docx",
        data=docx_bytes,
        file_name=f"{acronym}_Risk_Register.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="dl_docx"
    )
    st.success("✅ Word report ready.")

st.markdown("---")
section_label("🌐 Interactive HTML Dashboard")
st.markdown(
    f"<p style='color:{muted};font-size:0.85rem'>"
    f"Exports all risk charts as a self-contained HTML dashboard — "
    f"embed it in your project website with an iframe.</p>", unsafe_allow_html=True)

if st.button("🖥️ Generate HTML Dashboard", type="primary", key="html_btn"):
    with st.spinner("Building dashboard…"):
        html = _build_html()
    st.download_button(
        f"📥 Download {acronym}_Risk_Dashboard.html",
        data=html.encode("utf-8"),
        file_name=f"{acronym}_Risk_Dashboard.html",
        mime="text/html", key="dl_html"
    )
    st.success("✅ HTML dashboard ready — open in browser or embed with an <iframe>.")
